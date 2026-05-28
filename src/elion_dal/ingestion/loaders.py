"""Извлечение текста из PDF/DOCX (для сидинга и тестов; OCR не нужен — файлы текстовые)."""

from __future__ import annotations

from pathlib import Path


def load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts).strip()


def load_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def load_document(path: Path) -> str:
    """Диспетчер по расширению. Имя вида '*.docx.pdf' трактуется как PDF."""
    name = path.name.lower()
    if name.endswith(".pdf"):
        return load_pdf(path)
    if name.endswith(".docx") or name.endswith(".doc"):
        return load_docx(path)
    raise ValueError(f"Неподдерживаемый формат: {path.name}")
